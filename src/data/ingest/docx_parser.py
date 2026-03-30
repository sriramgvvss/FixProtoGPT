"""
Module: src.data.ingest.docx_parser
=====================================

Extract FIX protocol specification data from Word (.docx) documents.

Uses ``python-docx`` to iterate paragraphs and tables, applying
regex heuristics similar to the PDF parser.

Coding Standards
----------------
- PEP 8  : Python Style Guide
- PEP 257 : Docstring Conventions — Google-style docstrings
- PEP 484 : Type Hints
- Google Python Style Guide

Author : FixProtoGPT Team
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import List, Optional

from src.data.ingest.base import (
    CanonicalSpec,
    SpecKind,
    SpecParser,
    register_parser,
)

logger = logging.getLogger(__name__)

# Reuse the same regex patterns as the PDF parser
_TAG_DEF_RE = re.compile(
    r"(?:Tag|Field)\s*[:=]?\s*(\d{1,5})\s*[-–—=:]\s*([A-Za-z]\w+)",
    re.IGNORECASE,
)
_MSG_TYPE_RE = re.compile(
    r"MsgType\s*=\s*([A-Za-z0-9]+)\s*\(?\s*([A-Za-z]\w+)\s*\)?",
    re.IGNORECASE,
)


@register_parser
class DOCXParser(SpecParser):
    """Extract FIX specs from Word (.docx) documents."""

    EXTENSIONS = {".docx"}

    def can_handle(self, path: Path) -> bool:
        """Return ``True`` for ``.docx`` files."""
        return path.suffix.lower() in self.EXTENSIONS

    def parse(self, path: Path) -> List[CanonicalSpec]:
        """Parse a DOCX file and extract FIX specification records.

        Args:
            path: Path to the ``.docx`` file.

        Returns:
            List of :class:`CanonicalSpec` records.

        Raises:
            ImportError: If ``python-docx`` is not installed.
            FileNotFoundError: If *path* does not exist.
        """
        try:
            import docx  # python-docx
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install it with: pip install python-docx"
            ) from exc

        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {path}")

        logger.info("Parsing DOCX spec: %s", path.name)
        specs: List[CanonicalSpec] = []

        doc = docx.Document(str(path))

        # ── Tables — structured field definitions ─────────────────
        for table_idx, table in enumerate(doc.tables):
            specs.extend(self._parse_table(table, table_idx, path.name))

        # ── Paragraphs — free text and headings ───────────────────
        current_heading = ""
        paragraph_buffer: List[str] = []

        for para in doc.paragraphs:
            style = (para.style.name or "").lower()

            if "heading" in style:
                # Flush previous section
                if paragraph_buffer:
                    specs.extend(
                        self._parse_paragraph_block(
                            "\n".join(paragraph_buffer),
                            current_heading,
                            path.name,
                        )
                    )
                    paragraph_buffer.clear()
                current_heading = para.text.strip()
            else:
                text = para.text.strip()
                if text:
                    paragraph_buffer.append(text)

        # Flush last section
        if paragraph_buffer:
            specs.extend(
                self._parse_paragraph_block(
                    "\n".join(paragraph_buffer),
                    current_heading,
                    path.name,
                )
            )

        logger.info(
            "DOCX %s → %d canonical records",
            path.name, len(specs),
        )
        return specs

    # ── Helpers ───────────────────────────────────────────────────

    def _parse_table(
        self,
        table,
        table_idx: int,
        source: str,
    ) -> List[CanonicalSpec]:
        """Interpret a ``docx.table.Table`` as field definitions."""
        rows = table.rows
        if len(rows) < 2:
            return []

        # Read header row
        header = [cell.text.strip().lower() for cell in rows[0].cells]
        col_map = self._map_columns(header)
        if "tag" not in col_map:
            return []

        results: List[CanonicalSpec] = []
        for row in rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            tag = self._safe_int(self._cell(cells, col_map.get("tag")))
            if tag is None:
                continue

            results.append(
                CanonicalSpec(
                    kind=SpecKind.FIELD,
                    tag=tag,
                    name=self._cell(cells, col_map.get("name")),
                    data_type=(
                        self._cell(cells, col_map.get("type")).upper()
                        or None
                    ),
                    required=self._parse_bool(
                        self._cell(cells, col_map.get("required"))
                    ),
                    description=self._cell(cells, col_map.get("description")),
                    source=f"docx:{source}",
                    meta={"table_index": table_idx},
                )
            )
        return results

    def _parse_paragraph_block(
        self,
        text: str,
        heading: str,
        source: str,
    ) -> List[CanonicalSpec]:
        """Extract specs from a block of paragraphs under a heading."""
        results: List[CanonicalSpec] = []

        for m in _MSG_TYPE_RE.finditer(text):
            results.append(
                CanonicalSpec(
                    kind=SpecKind.MESSAGE,
                    msg_type=m.group(1),
                    name=m.group(2),
                    source=f"docx:{source}",
                    meta={"heading": heading},
                )
            )

        for m in _TAG_DEF_RE.finditer(text):
            tag = self._safe_int(m.group(1))
            if tag is None:
                continue
            results.append(
                CanonicalSpec(
                    kind=SpecKind.FIELD,
                    tag=tag,
                    name=m.group(2),
                    source=f"docx:{source}",
                    meta={"heading": heading},
                )
            )

        # Store raw text if no structure found but content is meaningful
        if not results and len(text) > 100:
            results.append(
                CanonicalSpec(
                    kind=SpecKind.RAW_TEXT,
                    name=heading or "docx_section",
                    description=text[:4000],
                    source=f"docx:{source}",
                    meta={"heading": heading},
                )
            )

        return results

    # ── Utilities (shared logic with PDF parser) ──────────────────

    @staticmethod
    def _map_columns(header: List[str]) -> dict:
        """Map DOCX table header names to logical field roles (tag, name, type, etc.)."""
        mapping = {}
        for idx, col in enumerate(header):
            if "tag" in col:
                mapping["tag"] = idx
            elif "name" in col or "field" in col:
                mapping["name"] = idx
            elif "type" in col:
                mapping["type"] = idx
            elif "req" in col:
                mapping["required"] = idx
            elif "desc" in col or "comment" in col:
                mapping["description"] = idx
        return mapping

    @staticmethod
    def _cell(cells: list, idx: Optional[int]) -> str:
        """Safely extract a stripped string from a cell list by index."""
        if idx is None or idx >= len(cells):
            return ""
        return str(cells[idx]).strip() if cells[idx] else ""

    @staticmethod
    def _safe_int(s: str) -> Optional[int]:
        """Parse a string to int, returning None on failure."""
        try:
            return int(s)
        except (ValueError, TypeError):
            return None

    @staticmethod
    def _parse_bool(s: str) -> Optional[bool]:
        """Interpret common truthy strings (y, yes, required, true, 1) as booleans."""
        if not s:
            return None
        return s.strip().lower() in {"y", "yes", "required", "true", "1"}
